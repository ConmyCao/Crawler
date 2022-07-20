import requests
from bs4 import BeautifulSoup
from docx import Document
# from docx.shared import Inches


def split_text_by_img(html, imglist):  # 将html代码以图片list进行分割成块
    content_parts = []  # 根据图标签将正文分割成N部分
    for imgtag in imglist:  # imgtag属性是bs4.element.Tag 后面需要使用str()函数转换成string
        # print(imgtag)
        html = str(html)  # 强制转化为字符串方便split分割
        str_tmp = html.split(str(imgtag))[0]  # 取图片分割的前一个元素 加入 正文list部分
        content_parts.append(str_tmp)
        # print(len(arr))
        html = html.replace((str_tmp + str(imgtag)), '')  # 将正文第一部分及图片标签字符串 从html中替换抹掉作为下一个for循环的html
        # print(html)
    content_parts.append(html)  # 把最后一张图片后的html内容补上
    return content_parts


def pic_down(referer_url, pic_url):  # 根据图片url保存图片，填写referer可伪装referer来源下载防盗链图片
    headers = {"Accept": "application/json",
               "Accept-Encoding": "gzip, deflate, br",
               "Accept-Language": "zh-CN,zh;q=0.9",
               "Connection": "keep-alive",
               'Cookie': 'Hm_lvt_0c0e9d9b1e7d617b3e6842e85b9fb068=1644906919,1644911381,1645167543,1646619975; '
                         'CNZZDATA1279807957=1568280098-1634697102-https%253A%252F%252Fwww.baidu.com%252F'
                         '%7C1646610516; locale=zh-CN; read_mode=day; default_font=font2; '
                         '_m7e_session_core=c813f019cec18a2a8f43c65eb9752864; '
                         'sensorsdata2015jssdkcross=%7B%22distinct_id%22%3A%2217c9bad52f36ad-0bc6d081956867-b7a1438'
                         '-2073600-17c9bad52f414dd%22%2C%22first_id%22%3A%22%22%2C%22props%22%3A%7B%22'
                         '%24latest_traffic_source_type%22%3A%22%E8%87%AA%E7%84%B6%E6%90%9C%E7%B4%A2%E6%B5%81%E9%87'
                         '%8F%22%2C%22%24latest_search_keyword%22%3A%22BeautifulSoup%E6%8A%93%E5%8F%96%E7%AE%80%E4%B9'
                         '%A6%E6%96%87%E7%AB%A0%22%2C%22%24latest_referrer%22%3A%22https%3A%2F%2Fwww.baidu.com%2Fs%22'
                         '%2C%22%24latest_utm_source%22%3A%22desktop%22%2C%22%24latest_utm_medium%22%3A%22search'
                         '-input%22%7D%2C%22%24device_id%22%3A%2217c9bad52f36ad-0bc6d081956867-b7a1438-2073600'
                         '-17c9bad52f414dd%22%7D; __yadk_uid=dyfYFtCXozcUPGdE849YvzC4hz4rvfgT; '
                         'ssxmod_itna2=WqGOY5YKGKBIIDl8p+oG=qDtDOnyIhP4GOUjbYoD8kaeGXwiGa8BBITUkbY8tz1OKGFbQTuLxe'
                         '=4e5h8XFjCK0+uzI7DvHqoreH+f36UGLHLY5+TZruskNW8HBv2pzjnfmnAmwrzkZnTmB94BQDH4rRx0Eeobhn=K'
                         '=Wquj0x8W9+e6HkmWOHNc0dflpb=1F48Da7dgPbqi47glDxXYBQeFWsqRa4aR2Tah2QyDyXPkvXLR2LCPC2z'
                         '=SiloWHDLrqe+MF6tnDa+A8v6GHCniOCv7UFRpW3NvrEgWt9R9THBI0NNcnROW4bDHKwre5xE+MU4WAm4TTpz'
                         '+qkBDurwA78m+lIAu/IpiBie+WIIieb8YoCFP3gRfSh5LedTu3PPHCoevYwUBrATtTwHWRlIhNgvHKF2GRTW9Tj'
                         '+tDbbu7TN9kKcTo3kmSWShXW7/zSXUcWi4HySa6fbrnTodTlfaBiSUf76486n9ABX3u8kur'
                         '=3PO8rvIri3PgGO67Fpy3D07ZB5miam4s=7aVDd/BDQKoQzKGYR5i7mDDLxD2YGDD===; '
                         'acw_tc=0b62601a16582990382593488e0140fcda2b9d08767994270417b9e0ee89fd; '
                         'acw_sc__v3=62d7a2a8d38e38dfc6872c8f6cef1442d91290e3; '
                         'ssxmod_itna=QqGxBDcmD=K6DXFG7GTG=3GQGCBNGkiiYjjDBk4x4iNDnD8x7YDvGGk3cYnY'
                         '=DNr8GPtMb1Gi3idDf3gmx8b=2GoDHxY=DUPc+KYD4+KGwD0eG+DD4DWDmmFDnxAQDjxGpnXvTs=DEDmb8DWPDYxDrE'
                         '=KDRxi7DDvd7x07DQH8GWi8K+p+RooCq7BeTmKD9roDshDfemIwEU3xBO7Aw'
                         '/yowx0kg40OnoHz8ooDU0IzcZ5e8j0odqhtahi3qiP5di0DHe2eG70Awf2PpQD5wi0P7ihop0UDD=; '
                         'signin_redirect=https://www.jianshu.com/p/2a1a7bbddce2?u_atoken=3e2558ca-9452-4894-a8b7'
                         '-a796f3c93f9a&u_asession=01xV6LrMyPi5sEoyRyaouY6U2exO6_Evm2QJstUPgsxs4bebXjMS2_kC8fKOD'
                         '-R6SGX0KNBwm7Lovlpxjd_P_q4JsKWYrT3W_NKPr8w6oU7K_cSuHVQtbOL4Krchkj85DbPpcarp92QKzyJKyYjREPlmBkFo3NEHBv0PZUm6pbxQU&u_asig=05QQGENPHj3y0qRSj9RI0XfgnnSf2HTCTSeb1eod9hofH1UkEwD35fFf1uZ9W6hGjMexBMVS2GSlIrOLMzhBuE8ExNb4iHCCcZsdSbrScneIa6kS_gpeW-2ip3IXSWaze49I5IdT05Vx2sYEm-rY9ZWKQNqoYolXXWwmhLMyJWW-X9JS7q8ZD7Xtz2Ly-b0kmuyAKRFSVJkkdwVUnyHAIJzQeR8AHMhwbE2htisfx0_JvuJu-vWw2_ZtEUEF_P514wWPRPQyB_SKrj-61LB_f61u3h9VXwMyh6PgyDIVSG1W-xH8762_l4mJ6MpQ5IMW1gfm5fsg5wmy8bWufzsZqoZGXsLmv5VkDonr9WzhLJLi5-_JfmYum_rcebAukmzt23mWspDxyAEEo4kbsryBKb9Q&u_aref=odETYQLEuBA2WfcEHIMpuatngtk%253D',
               "Host": "www.jianshu.com",
               "If-None-Match": 'W/"900805197f5e572f4a4bf253ed081267"',
               "sec-ch-ua": '".Not/A)Brand";v="99", "Google Chrome";v="103", "Chromium";v="103"',
               "sec-ch-ua-mobile": "?0",
               "sec-ch-ua-platform:": "Windows",
               "Sec-Fetch-Dest": "empty",
               "Sec-Fetch-Mode": "cors",
               "Sec-Fetch-Site": "same-origin",
               "Referer": referer_url,
               "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36",
               }
    # 保存图片至本地，因为新浪图片url中，不带后缀，这里就加了jpg后缀名，否则生成的word会报错
    img_name = pic_url.split('/')[-1] + '.jpg'
    # img_name = pic_url.split('/')[-1]

    with open(img_name, 'wb') as f:
        response = requests.get(pic_url, headers=headers).content
        f.write(response)
        f.close()
        return img_name


url = 'https://www.jianshu.com/p/2a1a7bbddce2'
headers_X = {"Accept": "application/json",
           "Accept-Encoding": "gzip, deflate, br",
           "Accept-Language": "zh-CN,zh;q=0.9",
           "Connection": "keep-alive",
           'Cookie': 'Hm_lvt_0c0e9d9b1e7d617b3e6842e85b9fb068=1644906919,1644911381,1645167543,1646619975; '
                     'CNZZDATA1279807957=1568280098-1634697102-https%253A%252F%252Fwww.baidu.com%252F'
                     '%7C1646610516; locale=zh-CN; read_mode=day; default_font=font2; '
                     '_m7e_session_core=c813f019cec18a2a8f43c65eb9752864; '
                     'sensorsdata2015jssdkcross=%7B%22distinct_id%22%3A%2217c9bad52f36ad-0bc6d081956867-b7a1438'
                     '-2073600-17c9bad52f414dd%22%2C%22first_id%22%3A%22%22%2C%22props%22%3A%7B%22'
                     '%24latest_traffic_source_type%22%3A%22%E8%87%AA%E7%84%B6%E6%90%9C%E7%B4%A2%E6%B5%81%E9%87'
                     '%8F%22%2C%22%24latest_search_keyword%22%3A%22BeautifulSoup%E6%8A%93%E5%8F%96%E7%AE%80%E4%B9'
                     '%A6%E6%96%87%E7%AB%A0%22%2C%22%24latest_referrer%22%3A%22https%3A%2F%2Fwww.baidu.com%2Fs%22'
                     '%2C%22%24latest_utm_source%22%3A%22desktop%22%2C%22%24latest_utm_medium%22%3A%22search'
                     '-input%22%7D%2C%22%24device_id%22%3A%2217c9bad52f36ad-0bc6d081956867-b7a1438-2073600'
                     '-17c9bad52f414dd%22%7D; __yadk_uid=dyfYFtCXozcUPGdE849YvzC4hz4rvfgT; '
                     'ssxmod_itna2=WqGOY5YKGKBIIDl8p+oG=qDtDOnyIhP4GOUjbYoD8kaeGXwiGa8BBITUkbY8tz1OKGFbQTuLxe'
                     '=4e5h8XFjCK0+uzI7DvHqoreH+f36UGLHLY5+TZruskNW8HBv2pzjnfmnAmwrzkZnTmB94BQDH4rRx0Eeobhn=K'
                     '=Wquj0x8W9+e6HkmWOHNc0dflpb=1F48Da7dgPbqi47glDxXYBQeFWsqRa4aR2Tah2QyDyXPkvXLR2LCPC2z'
                     '=SiloWHDLrqe+MF6tnDa+A8v6GHCniOCv7UFRpW3NvrEgWt9R9THBI0NNcnROW4bDHKwre5xE+MU4WAm4TTpz'
                     '+qkBDurwA78m+lIAu/IpiBie+WIIieb8YoCFP3gRfSh5LedTu3PPHCoevYwUBrATtTwHWRlIhNgvHKF2GRTW9Tj'
                     '+tDbbu7TN9kKcTo3kmSWShXW7/zSXUcWi4HySa6fbrnTodTlfaBiSUf76486n9ABX3u8kur'
                     '=3PO8rvIri3PgGO67Fpy3D07ZB5miam4s=7aVDd/BDQKoQzKGYR5i7mDDLxD2YGDD===; '
                     'acw_tc=0b62601a16582990382593488e0140fcda2b9d08767994270417b9e0ee89fd; '
                     'acw_sc__v3=62d7a2a8d38e38dfc6872c8f6cef1442d91290e3; '
                     'ssxmod_itna=QqGxBDcmD=K6DXFG7GTG=3GQGCBNGkiiYjjDBk4x4iNDnD8x7YDvGGk3cYnY'
                     '=DNr8GPtMb1Gi3idDf3gmx8b=2GoDHxY=DUPc+KYD4+KGwD0eG+DD4DWDmmFDnxAQDjxGpnXvTs=DEDmb8DWPDYxDrE'
                     '=KDRxi7DDvd7x07DQH8GWi8K+p+RooCq7BeTmKD9roDshDfemIwEU3xBO7Aw'
                     '/yowx0kg40OnoHz8ooDU0IzcZ5e8j0odqhtahi3qiP5di0DHe2eG70Awf2PpQD5wi0P7ihop0UDD=; '
                     'signin_redirect=https://www.jianshu.com/p/2a1a7bbddce2?u_atoken=3e2558ca-9452-4894-a8b7'
                     '-a796f3c93f9a&u_asession=01xV6LrMyPi5sEoyRyaouY6U2exO6_Evm2QJstUPgsxs4bebXjMS2_kC8fKOD'
                     '-R6SGX0KNBwm7Lovlpxjd_P_q4JsKWYrT3W_NKPr8w6oU7K_cSuHVQtbOL4Krchkj85DbPpcarp92QKzyJKyYjREPlmBkFo3NEHBv0PZUm6pbxQU&u_asig=05QQGENPHj3y0qRSj9RI0XfgnnSf2HTCTSeb1eod9hofH1UkEwD35fFf1uZ9W6hGjMexBMVS2GSlIrOLMzhBuE8ExNb4iHCCcZsdSbrScneIa6kS_gpeW-2ip3IXSWaze49I5IdT05Vx2sYEm-rY9ZWKQNqoYolXXWwmhLMyJWW-X9JS7q8ZD7Xtz2Ly-b0kmuyAKRFSVJkkdwVUnyHAIJzQeR8AHMhwbE2htisfx0_JvuJu-vWw2_ZtEUEF_P514wWPRPQyB_SKrj-61LB_f61u3h9VXwMyh6PgyDIVSG1W-xH8762_l4mJ6MpQ5IMW1gfm5fsg5wmy8bWufzsZqoZGXsLmv5VkDonr9WzhLJLi5-_JfmYum_rcebAukmzt23mWspDxyAEEo4kbsryBKb9Q&u_aref=odETYQLEuBA2WfcEHIMpuatngtk%253D',
           "Host": "www.jianshu.com",
           "If-None-Match": 'W/"900805197f5e572f4a4bf253ed081267"',
           "sec-ch-ua": '".Not/A)Brand";v="99", "Google Chrome";v="103", "Chromium";v="103"',
           "sec-ch-ua-mobile": "?0",
           # "sec-ch-ua-platform:": "Windows",
           "Sec-Fetch-Dest": "empty",
           "Sec-Fetch-Mode": "cors",
           "Sec-Fetch-Site": "same-origin",
           "Referer": url,
           "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36",
           }
html = requests.get(url, headers=headers_X).content
print("the html is"+str(html))
soup = BeautifulSoup(html, 'html.parser')

# title= soup.select('section>h1').string
title = soup.title.string  # 取得浏览器标签的文字内容
print('the title is:', str(title))
# print('the title2 is:', str(title2))

post_detial = soup.select('article')

print(str(post_detial))
# # img_tag_list=post_detial.select('a img')
# img_tag_list = post_detial.select('div > img')  # 在详情内容中截取出图片img的标签list
img_tag_list = soup.select('article div>img')
print("发现图片标签个数：" + str(len(img_tag_list)))
new_text = split_text_by_img(post_detial, img_tag_list)  # 按图片分割成段，返回list串
print("截取博客文章片段数：" + str(len(new_text)))

# # print('beg==============================================')
# # print(new_text[1])
# # print('end==============================================')

document = Document()
document.add_heading(title)  # 向文档里添加标题
i = 0
for part in new_text:  # 循环写入
    print('写入第 ' + str(i + 1) + ' 个片段')
    # print(n)
    part = '<html><body><div >' + part  # part是含html标签的字符串，下面使用BeautifuSoup时需要lxml格式化，所以需要加前缀
    # 使的每个part部分都更像一个网页，否则BeautifulSoup(part,'lxml')处理的时候会把第二部分开始的内容处理为空

    part_tag = BeautifulSoup(part, 'lxml')  # 如果不进行lxml处理，下面get_text()无法使用。
    document.add_paragraph(part_tag.get_text())  # 向文档里添加文字

    if (i < len(img_tag_list)):  # 写完片段后紧接写入位于此处分割的图片
        imgurl = img_tag_list[i].get('data-src')  # 新浪图片地址在real_src属性里，一般是src
        img_name = pic_down('https://jianshu.com/', imgurl)  # 新浪图片有防盗链，需要加入referer='http://blog.sina.com.cn'
        print('写入第 ' + str(i + 1) + ' 张图片：' + imgurl)
        # document.add_paragraph('此处需要插入图片'+img_name+imgurl)#向文档里添加文字
        document.add_picture(img_name)  # 向文档里添加图片
        # os.remove(img_name)#删除保存在本地的图片
    i = i + 1

# document.save('图文.doc')#保存文档
document.save(title + '.doc')