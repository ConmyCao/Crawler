import requests
from bs4 import BeautifulSoup
from docx import Document
# from docx.shared import Inches


def split_text_by_img(html, imglist):  # 将html代码以图片list进行分割成块
    content_parts = []  # 根据图标签将正文分割成N部分
    for imgtag in imglist:  # imgtag属性是bs4.element.Tag 后面需要使用str()函数转换成string
        # print(imgtag)
        html = str(html)  # 强制转化为字符串方便split分割
        str_tmp = html.split(str(imgtag))[0]  # 取图片分割的前一个元素 加入 正文list部分
        content_parts.append(str_tmp)
        # print(len(arr))
        html = html.replace((str_tmp + str(imgtag)), '')  # 将正文第一部分及图片标签字符串 从html中替换抹掉作为下一个for循环的html
        # print(html)
    content_parts.append(html)  # 把最后一张图片后的html内容补上
    return content_parts


def pic_down(referer_url, pic_url):  # 根据图片url保存图片，填写referer可伪装referer来源下载防盗链图片
    headers = {"Accept": "/",
               "x-zse-93": "101_3_3.0",
               "x-zse-96": "2.0_qJh//4tVnPHeoD7fVC7RolJZ+I5aaq6Hf1VeF6bwK+NkOJf9SOsStNL5MPNp1PbG",
               "Accept-Encoding": "gzip, deflate, br",
               "Accept-Language": "zh-CN,zh;q=0.9",
               "Referer": referer_url,
               "User-Agent": "Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/42.0.2311.90 Safari/537.36",
               'Cookie': '_zap=11e68ed0-a0af-48ee-bccc-0356aa4d784c; d_c0="ACAehWXr5ROPThqVFNKb3gYeomnAMjKpNGM=|1634627152"; _9755xjdesxxd_=32; YD00517437729195%3AWM_TID=IdrCPNcJvChBABFRVFI%2BoNO12SXdwOuI; ISSW=1; YD00517437729195%3AWM_NIKE=9ca17ae2e6ffcda170e2e6eeb3d5408392fc9bee598a9a8ea3d85f939e9baef57e89869992d57df2b685a3ea2af0fea7c3b92ab29596d6c653f595adb8d245a6a6a590d879b5afad86f125b8b9bdd1d66392a7f783e459909ca2d4ae46869fe599bc25a688fdccc866819197acc45483b69ed8f152b5e9fbb7f94f95a683b8cf7bba8c9bb9bc3db3988b94db618cbeaaaaf35db8f0b9bbca73f4b5a086c85cf4f5a1d9c87ef4b7fbd6d466aeae869ad94faaa9ac8dee37e2a3; YD00517437729195%3AWM_NI=rKZ5GY7KcWSOfzkoUswqcer%2Ftgs0WiRy9baSLDn3%2FNFA%2Bx9%2FfX3%2BEnCOYKAXAjR7rZw75VzEWzMWLqzTUX97qHEsXkHbfIpWt0wZKy0kPeEP0%2Bo8Ev0i55QnnrJyZEjwdlo%3D; gdxidpyhxdE=vc14ndA9zLMXo46qVyZZcXNpiShqJmzWe0zy%2BGulMIxtjCIS8f2ZaLTsHHeCtKkOosEnKp%5CUnrvMxsSxooEzuLrHcgOuC3tbkfqZjsii%2FUi4%2BzzXjfYiKa7aCbq3GCkZGqUcc7LAIUmN%2BDx%2F1wBgD15cG4zZ6SM%2F2VxU15X%2F4%2FOw%5ChUW%3A1644202251262; z_c0="2|1:0|10:1644201801|4:z_c0|92:Mi4xS2pqOUFnQUFBQUFBSUI2RlpldmxFeVlBQUFCZ0FsVk5TZFh0WWdDb0VaRXAwellaSGMtbmprNmFxalBlLXg0WUhn|9ab5061b18947fb9fb97383703503d11602227da69ed0c8f48a6422a13dbf2d6"; Hm_lvt_98beee57fd2ef70ccdd5ca52b9740c49=1646614006,1647483891,1647506369,1647847078; q_c1=581c48fbabb448f18467f031d4791cc3|1653612367000|1653612367000; _xsrf=4541ce2b-6ad0-434d-8c2c-6afd5e533e33; NOT_UNREGISTER_WAITING=1; KLBRSID=ed2ad9934af8a1f80db52dcb08d13344|1657614737|1657614659',
               'x-ab-pb': 'CpQCGwA/AEMARwC0AEABaQFqAXQBOwJ9ArkCzALXAtgCMgNPA1ADoAOhA6IDtwPzA/QDMwSMBI0EpgTWBOkEEQUpBTIFUQWLBYwFngUWBjAGMQZBBn4GlAaiBusGJwdXB3cHeAebB9gH3AfdBycIZwh0CHYIeQjFCNYI2gjlCAEJFgk/CUIJSQlUCVUJYAmECY0JqwnDCcQJxQnGCccJyAnJCcoJywnMCdEJ5QnxCfQJ9gkECkkKZQprCoMKmAqlCqkKvgrECtQK3QrtCv0K/goTCykLOwtDC0YLcQt2C3kLfQuFC4cLjQujC7kLwAvXC+AL5QvmCywMMQw4DDQM3AtgCwELmwu1C+QK4AvsClILVgy0CjcMEooBAAAVAAABAQAAAAEVAAAAAAAAAAAAAAAAAAQEAAQAAAABAQAAAQAAAQAAAAQAAAEAAAAAAgIABAAABgAAAAEAAAAAAQAAAQAAAAAAAAAAAAAAAAAAAAAAAQAAAAAEAQAAAAEBAQAEAAELAAAAAAADAAAAAQEAAAEBAAABAQAAAAAAAgMAAAEBAQAB'
               }
    # 保存图片至本地，因为新浪图片url中，不带后缀，这里就加了jpg后缀名，否则生成的word会报错
    img_name = pic_url.split('/')[-1] + '.jpg'
    # img_name = pic_url.split('/')[-1]

    with open(img_name, 'wb') as f:
        response = requests.get(pic_url, headers=headers).content
        f.write(response)
        f.close()
        return img_name


url = 'https://ranjuan.cn/web-access-limit/'
html = requests.get(url).content
# print("the html is"+str(html))
soup = BeautifulSoup(html, 'html.parser')

title= soup.select('article > header > h3')[0].string
title2 = soup.title.string  # 取得浏览器标签的文字内容
print('the title is:', str(title))
print('the title2 is:', str(title2))

post_detial = soup.select('article>div')[0]

print(str(post_detial))
# # img_tag_list=post_detial.select('a img')
# img_tag_list = post_detial.select('div > img')  # 在详情内容中截取出图片img的标签list
img_tag_list = soup.select('article>div>figure>img')
print("发现图片标签个数：" + str(len(img_tag_list)))
new_text = split_text_by_img(post_detial, img_tag_list)  # 按图片分割成段，返回list串
print("截取博客文章片段数：" + str(len(new_text)))

# # print('beg==============================================')
# # print(new_text[1])
# # print('end==============================================')

document = Document()
document.add_heading(title)  # 向文档里添加标题
i = 0
for part in new_text:  # 循环写入
    print('写入第 ' + str(i + 1) + ' 个片段')
    # print(n)
    part = '<html><body><div >' + part  # part是含html标签的字符串，下面使用BeautifuSoup时需要lxml格式化，所以需要加前缀
    # 使的每个part部分都更像一个网页，否则BeautifulSoup(part,'lxml')处理的时候会把第二部分开始的内容处理为空

    part_tag = BeautifulSoup(part, 'lxml')  # 如果不进行lxml处理，下面get_text()无法使用。
    document.add_paragraph(part_tag.get_text())  # 向文档里添加文字

    if (i < len(img_tag_list)):  # 写完片段后紧接写入位于此处分割的图片
        imgurl = img_tag_list[i].get('data-src')  # 新浪图片地址在real_src属性里，一般是src
        img_name = pic_down('https://ranjuan.cn/', imgurl)  # 新浪图片有防盗链，需要加入referer='http://blog.sina.com.cn'
        print('写入第 ' + str(i + 1) + ' 张图片：' + imgurl)
        # document.add_paragraph('此处需要插入图片'+img_name+imgurl)#向文档里添加文字
        document.add_picture(img_name)  # 向文档里添加图片
        # os.remove(img_name)#删除保存在本地的图片
    i = i + 1

# document.save('图文.doc')#保存文档
document.save(title + '.doc')